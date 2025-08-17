import database
import config
import os
try:
    from docx import Document
    from docx.shared import Inches
    import PyPDF2
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx and PyPDF2 not available. Word and PDF content extraction disabled.")

class UserInputError(Exception):
    """Исключение при некорректно введённых данных."""

class Process:
    """Базовый класс для многошаговых процессов с поддержкой отмены."""
    CANCEL_KEYWORD = "отмена"
    _bot = None

    def __init__(self, ID, cancelable = True):
        self._info = database.Client(ID)
        self._chain = []
        self._max_i = len(self._chain) - 1
        self._i = 0
        self._is_active = False
        self._cancelable = cancelable

        self._current_request = ""
        self._canceled = False

    def update_last_request(self, request):
        self._current_request = request

        if self._cancelable:
            if self._current_request == self.CANCEL_KEYWORD:
                # помечаем как отменённый, останавливаем и уведомляем пользователя
                self._canceled = True
                self.stop()
                try:
                    if self._bot and self._info:
                        self._bot.send_message(self._info.get_ID(), "Действие отменено")
                except Exception:
                    pass

    def stop(self):
        self._i = 0
        self._is_active = False

    def out(self, text: str, keyboard = None):
        try:
            if self._bot and self._info:
                self._bot.send_message(self._info.get_ID(), text, reply_markup = keyboard)
        except Exception as e:
            try:
                print(f"Не удалось отправить сообщение из процесса: {e}")
            except Exception:
                pass

    def execute(self):
        try:
            # если предыдущий ввод был отменой — ничего не делать
            if getattr(self, "_canceled", False):
                self._canceled = False
                return
            # помечаем процесс как активный на время выполнения шага
            self._is_active = True
            # Выполняем текущий шаг
            self._chain[self._i]()
            # Переходим к следующему шагу или завершаем
            if self._i < self._max_i:
                self._i += 1
            else:
                self.stop()
                return

            # Если следующий шаг — это ask_* (запрос ввода), выполняем его сразу,
            # чтобы пользователь сразу получил следующий вопрос.
            if self._is_active and self._i <= self._max_i:
                try:
                    next_step = self._chain[self._i]
                    step_name = getattr(next_step, "__name__", "")
                    if isinstance(step_name, str) and step_name.startswith("ask_"):
                        next_step()
                        if self._i < self._max_i:
                            self._i += 1
                        else:
                            self.stop()
                except Exception as inner_e:
                    print(f"Ошибка при выполнении следующего шага процесса: {inner_e}")
        except UserInputError as e:
            print(f"Пользователь {self._info.name} неверно ввёл запрошенные данные")

    @classmethod
    def set_bot(cls, bot_instance):
        cls._bot = bot_instance


class FileSender:
    IMAGE_EXTENSIONS = ["png", "jpg", "jpeg", "webp"]
    DOCUMENT_EXTENSIONS = ["doc", "docx", "pdf", "xlsx", "xls", "ppt", "pptx", "txt"]
    AUDIO_EXTENSIONS = ["mp3", "wav", "ogg", "m4a", "aac", "flac"]
    VIDEO_EXTENSIONS = ["mp4", "avi", "mkv", "mov", "wmv", "flv", "webm"]

    unzipped_text_document = True # позволить боту отправить сразу текст из файла
    bot = None
    chat_id = None
    
    # Базовая директория проекта (на уровень выше папки program)
    PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

    def __init__(self, keyboard=None, *args):
        self.keyboard = keyboard

        self.file_handlers = {}

        for path in args:
            extension = file_extension(path)
            if extension in self.IMAGE_EXTENSIONS:
                self.file_handlers[path] = self.__push_image
            elif extension in self.DOCUMENT_EXTENSIONS:
                # Извлекаем содержимое документов как текст
                if self.unzipped_text_document:
                    self.file_handlers[path] = self.__push_document_content
                else:
                    self.file_handlers[path] = self.__push_document
            elif extension in self.AUDIO_EXTENSIONS:
                self.file_handlers[path] = self.__push_audio
            elif extension in self.VIDEO_EXTENSIONS:
                self.file_handlers[path] = self.__push_video

    @classmethod
    def set_chat_id(cls, chat_id):
        cls.chat_id = chat_id

    @classmethod
    def set_bot(cls, bot_instance):
        cls.bot = bot_instance

    def push(self):
        sent = 0
        required_to_send = len(self.file_handlers.keys())
        for path in self.file_handlers.keys():
            sent += 1
            if sent == required_to_send:
                self.file_handlers[path](path, keyboard = None)
            else:
                self.file_handlers[path](path, keyboard = self.keyboard)

    def _resolve_path(self, path: str) -> str:
        """Разрешает путь к файлу относительно cwd или корня проекта."""
        try:
            if os.path.isabs(path):
                return path
            if os.path.exists(path):
                return path
            candidate = os.path.normpath(os.path.join(self.PROJECT_ROOT, path))
            return candidate
        except Exception as e:
            print(f"Ошибка разрешения пути '{path}': {e}")
            return path

    def __push_image(self, path: str, keyboard = None, caption: str = None):
        resolved = self._resolve_path(path)
        with open(resolved, 'rb') as file:
            self.bot.send_photo(self.chat_id, file, reply_markup = keyboard, caption=caption)

    def __push_document(self, path: str, keyboard = None, caption: str = None):
        resolved = self._resolve_path(path)
        with open(resolved, 'rb') as file:
            self.bot.send_document(self.chat_id, file, reply_markup = keyboard, caption=caption)

    def __push_audio(self, path: str, keyboard = None, caption: str = None):
        resolved = self._resolve_path(path)
        with open(resolved, 'rb') as file:
            self.bot.send_audio(self.chat_id, file, reply_markup = keyboard, caption=caption)

    def __push_video(self, path: str, keyboard = None, caption: str = None):
        resolved = self._resolve_path(path)
        with open(resolved, 'rb') as file:
            self.bot.send_video(self.chat_id, file, reply_markup = keyboard, caption=caption)

    def __push_unzipped_text_document(self, path: str, keyboard = None, caption: str = None):
        resolved = self._resolve_path(path)
        with open(resolved, 'r', encoding='utf-8') as file:
            text = file.read()

        self.bot.send_message(self.chat_id, f"{text}", reply_markup = keyboard)
    
    def __extract_docx_content(self, path: str) -> str:
        """Извлекает текст из Word документа."""
        try:
            if not DOCX_AVAILABLE:
                return "Ошибка: библиотека python-docx не установлена. Установите: pip install python-docx"
            
            resolved = self._resolve_path(path)
            doc = Document(resolved)
            
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            # Добавляем информацию о таблицах, если есть
            if doc.tables:
                content_parts.append("\n📊 Документ содержит таблицы с дополнительной информацией.")
            
            return "\n\n".join(content_parts) if content_parts else "Документ не содержит текста."
            
        except Exception as e:
            return f"Ошибка при чтении Word документа: {e}"
    
    def __extract_pdf_content(self, path: str) -> str:
        """Извлекает текст из PDF документа."""
        try:
            if not DOCX_AVAILABLE:
                return "Ошибка: библиотека PyPDF2 не установлена. Установите: pip install PyPDF2"
            
            resolved = self._resolve_path(path)
            content_parts = []
            
            with open(resolved, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text().strip()
                    if page_text:
                        content_parts.append(f"📄 Страница {page_num}:\n{page_text}")
            
            return "\n\n".join(content_parts) if content_parts else "PDF документ не содержит извлекаемого текста."
            
        except Exception as e:
            return f"Ошибка при чтении PDF документа: {e}"
    
    def __push_document_content(self, path: str, keyboard = None, caption: str = None):
        """Отправляет содержимое документа как текст."""
        try:
            extension = file_extension(path).lower()
            
            if extension == "docx":
                content = self.__extract_docx_content(path)
            elif extension == "pdf":
                content = self.__extract_pdf_content(path)
            elif extension == "txt":
                resolved = self._resolve_path(path)
                with open(resolved, 'r', encoding='utf-8') as file:
                    content = file.read()
            else:
                # Для других типов документов отправляем как файл
                self.__push_document(path, keyboard, caption)
                return
            
            # Разбиваем длинный текст на части (Telegram лимит ~4096 символов)
            max_length = 4000
            if len(content) <= max_length:
                self.bot.send_message(self.chat_id, content, reply_markup=keyboard)
            else:
                # Разбиваем на части
                parts = []
                current_part = ""
                
                for line in content.split('\n'):
                    if len(current_part) + len(line) + 1 <= max_length:
                        current_part += line + '\n'
                    else:
                        if current_part:
                            parts.append(current_part.strip())
                        current_part = line + '\n'
                
                if current_part:
                    parts.append(current_part.strip())
                
                # Отправляем части
                for i, part in enumerate(parts):
                    if i == len(parts) - 1:  # Последняя часть с клавиатурой
                        self.bot.send_message(self.chat_id, f"📄 Часть {i+1}/{len(parts)}:\n\n{part}", reply_markup=keyboard)
                    else:
                        self.bot.send_message(self.chat_id, f"📄 Часть {i+1}/{len(parts)}:\n\n{part}")
                        
        except Exception as e:
            print(f"Ошибка при отправке содержимого документа {path}: {e}")
            # Fallback: отправляем как файл
            self.__push_document(path, keyboard, caption)


class Validator:
    """Статические валидаторы пользовательского ввода (русский язык)."""
    RU_ALPHABET = "абвгдеёжзийклмнопрстуфхцчшщъыьэюя"
    RU_CLASS_PARALLEL = "абвг"
    DECIMAL_DIGITS = "0123456789"

    @staticmethod
    def name(string: str, alphabet: str = RU_ALPHABET):
        """Разрешает только русские буквы для имени."""
        try:
            string = string.lower()
            for char in string:
                if char not in alphabet:
                    return False
            return True
        except Exception as e:
            print(f"Ошибка при валидации имени: {e}")
            return False
    
    @classmethod
    def surname(cls, string: str, lang: str = RU_ALPHABET):
        return cls.name(string, lang)

    @staticmethod
    def city(string: str, lang: str = RU_ALPHABET):
        """Разрешает русские буквы и пробелы в названии города."""
        try:
            string = string.lower()
            return all(char in lang or char == " " for char in string)
        except Exception as e:
            print(f"Ошибка при валидации города: {e}")
            return False

    @staticmethod
    def class_number(number: str, parallel: str = RU_CLASS_PARALLEL):
        """Проверяет формат класса (1..11 + буква из `parallel`)."""
        try:
            class_num = int(number[:-1])
            letter = number[-1].lower()
            return 1 <= class_num <= 11 and letter in parallel
        except (ValueError, IndexError) as e:
            print(f"Ошибка при валидации номера класса: {e}")
            return False

    @staticmethod
    def school(number: str):
        """Извлекает номер школы в конце строки и проверяет, что он > 0."""
        try:
            extract = number.split(" ")[-1]
            return int(extract) > 0
        except (ValueError, IndexError) as e:
            print(f"Ошибка при валидации номера школы: {e}")
            return False

    @classmethod
    def create_password(cls, string: str, max_len: int = config.PASSWORD_LENGTH):
        """Проверяет, что пароль цифровой и длиной `max_len`."""
        try:
            if len(string) != max_len:
                return False
            return all(numeral in cls.DECIMAL_DIGITS for numeral in string)
        except Exception as e:
            print(f"Ошибка при валидации пароля: {e}")
            return False


def is_cancel(self) -> bool:
    """Простой хелпер отмены: возвращает True, если введено слово отмены."""
    try:
        last_request = getattr(self, "_current_request", None)
        return isinstance(last_request, str) and last_request == Process.CANCEL_KEYWORD
    except Exception:
        return False


def extractor(collection: list[dict], key) -> list:
    """Извлекает значения "key" из списка словарей."""
    extracted = []

    for dictonary in collection:
        find = dictonary.get(key)
        extracted.append(find)
    
    return extracted

def file_extension(path: str) -> str:
    """Определяет расширение файла."""
    return path.split(".")[-1]

def transform_request(request: str):
    """Преобразование строки в нижний регистр и замена "ё" на "е"."""
    request = request.lower().strip()
    if "ё" in request:
        request = request.replace("ё", "е")
    return request